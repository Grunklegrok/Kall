import hashlib
import json
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from sqlmodel import Session, select

from kall.models import (
    CoverLetterChange,
    CoverLetterProposal,
    DocumentArtifact,
    DocumentGenerationAudit,
    GeneratedDocument,
    Job,
    JobRequirementAnalysis,
    KeywordCoverageReport,
    TailoringChange,
    TailoringProposal,
)

GENERATION_VERSION = "documents-v1"


def _sha(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _accepted_text(change: TailoringChange) -> str | None:
    if change.status == "accepted":
        return change.proposed_text
    if change.status == "edited":
        return change.edited_text
    return None


def finalized_resume_content(session: Session, proposal: TailoringProposal) -> list[dict[str, str]]:
    if proposal.status != "finalized" or not proposal.finalized_at:
        raise ValueError("Tailoring proposal must be finalized before document generation")
    changes = list(
        session.exec(
            select(TailoringChange)
            .where(TailoringChange.proposal_id == proposal.id)
            .order_by(TailoringChange.id)
        )
    )
    if any(change.status == "pending" for change in changes):
        raise ValueError("Every tailoring change must be reviewed")
    return [
        {"section": change.section, "text": text}
        for change in changes
        if (text := _accepted_text(change))
    ]


def keyword_report(
    session: Session,
    generated: GeneratedDocument,
    job_id: int,
    content: str,
) -> KeywordCoverageReport:
    analysis = session.exec(
        select(JobRequirementAnalysis).where(JobRequirementAnalysis.job_id == job_id)
    ).first()
    required = analysis.required_skills if analysis else []
    preferred = analysis.preferred_skills if analysis else []
    lowered = content.lower()
    required_covered = [term for term in required if term.lower() in lowered]
    preferred_covered = [term for term in preferred if term.lower() in lowered]
    unsupported = [term for term in required + preferred if term.lower() not in lowered]
    return KeywordCoverageReport(
        generated_document_id=generated.id,
        required_covered=required_covered,
        preferred_covered=preferred_covered,
        unsupported=list(dict.fromkeys(unsupported)),
        required_percent=round(100 * len(required_covered) / len(required)) if required else 100,
        preferred_percent=round(100 * len(preferred_covered) / len(preferred)) if preferred else 100,
    )


def _write_docx(path: Path, title: str, sections: Iterable[dict[str, str]]) -> bytes:
    document = Document()
    document.add_heading(title, 0)
    for item in sections:
        document.add_heading(item["section"].replace("_", " ").title(), level=1)
        document.add_paragraph(item["text"])
    document.save(path)
    return path.read_bytes()


def _write_pdf(path: Path, title: str, sections: Iterable[dict[str, str]]) -> bytes:
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 18)]
    for item in sections:
        story.extend(
            [
                Paragraph(item["section"].replace("_", " ").title(), styles["Heading2"]),
                Paragraph(item["text"], styles["BodyText"]),
                Spacer(1, 12),
            ]
        )
    SimpleDocTemplate(str(path), pagesize=LETTER, title=title).build(story)
    return path.read_bytes()


def generate_resume_documents(
    session: Session,
    proposal: TailoringProposal,
    template_key: str = "standard",
) -> GeneratedDocument:
    sections = finalized_resume_content(session, proposal)
    content_text = "\n\n".join(item["text"] for item in sections)
    canonical = json.dumps(sections, sort_keys=True).encode()
    generated = GeneratedDocument(
        user_id=proposal.user_id,
        proposal_id=proposal.id,
        job_id=proposal.job_id,
        resume_id=proposal.resume_id,
        document_type="resume",
        template_key=template_key,
        content_json={"sections": sections},
        checksum=_sha(canonical),
        finalized_at=datetime.utcnow(),
    )
    session.add(generated)
    session.commit()
    session.refresh(generated)

    output = Path("generated") / str(proposal.user_id) / f"document-{generated.id}"
    output.mkdir(parents=True, exist_ok=True)
    artifacts: list[tuple[str, str, bytes]] = []
    txt = content_text.encode("utf-8")
    txt_path = output / "resume.txt"
    txt_path.write_bytes(txt)
    artifacts.append(("txt", "text/plain", txt))
    artifacts.append(("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", _write_docx(output / "resume.docx", "Tailored Resume", sections)))
    artifacts.append(("pdf", "application/pdf", _write_pdf(output / "resume.pdf", "Tailored Resume", sections)))

    for extension, mime, data in artifacts:
        session.add(
            DocumentArtifact(
                generated_document_id=generated.id,
                format=extension,
                file_path=str(output / f"resume.{extension}"),
                mime_type=mime,
                byte_size=len(data),
                checksum=_sha(data),
            )
        )
    report = keyword_report(session, generated, proposal.job_id, content_text)
    session.add(report)
    session.add(
        DocumentGenerationAudit(
            user_id=proposal.user_id,
            generated_document_id=generated.id,
            event="resume_generated",
            details={"template": template_key, "generation_version": GENERATION_VERSION},
        )
    )
    session.commit()
    return generated


def propose_cover_letter(
    session: Session,
    proposal: TailoringProposal,
    emphasis: str,
    tone: str,
    length: str,
    company_interest_notes: str | None,
) -> CoverLetterProposal:
    if proposal.status != "finalized":
        raise ValueError("Finalize the resume tailoring proposal first")
    job = session.get(Job, proposal.job_id)
    sections = finalized_resume_content(session, proposal)
    evidence = [item["text"] for item in sections if item["section"] == "achievement"][:2]
    letter = CoverLetterProposal(
        user_id=proposal.user_id,
        job_id=proposal.job_id,
        tailoring_proposal_id=proposal.id,
        emphasis=emphasis,
        tone=tone,
        length=length,
        company_interest_notes=company_interest_notes,
    )
    session.add(letter)
    session.commit()
    session.refresh(letter)
    paragraphs = [
        f"I am applying for the {job.title if job else 'position'} at {job.company if job else 'your organization'}.",
        " ".join(evidence) if evidence else "My verified experience aligns with the responsibilities described in the role.",
        company_interest_notes or "I would welcome the opportunity to discuss how my experience can support the team.",
    ]
    for position, text in enumerate(paragraphs):
        session.add(
            CoverLetterChange(
                proposal_id=letter.id,
                position=position,
                proposed_text=text,
                evidence=[{"source": "finalized_resume", "proposal_id": proposal.id}],
            )
        )
    session.commit()
    return letter


def review_cover_letter_change(
    session: Session,
    change: CoverLetterChange,
    decision: str,
    edited_text: str | None = None,
) -> CoverLetterChange:
    if decision not in {"accepted", "edited", "rejected"}:
        raise ValueError("Invalid decision")
    if decision == "edited" and not edited_text:
        raise ValueError("Edited text is required")
    change.status = decision
    change.edited_text = edited_text if decision == "edited" else None
    change.reviewed_at = datetime.utcnow()
    session.add(change)
    session.commit()
    session.refresh(change)
    return change


def finalize_cover_letter(session: Session, proposal: CoverLetterProposal) -> CoverLetterProposal:
    changes = list(session.exec(select(CoverLetterChange).where(CoverLetterChange.proposal_id == proposal.id)))
    if not changes or any(change.status == "pending" for change in changes):
        raise ValueError("Every cover letter paragraph must be reviewed")
    proposal.status = "finalized"
    proposal.finalized_at = datetime.utcnow()
    session.add(proposal)
    session.commit()
    session.refresh(proposal)
    return proposal
